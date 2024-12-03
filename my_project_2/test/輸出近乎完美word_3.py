from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import json

def add_custom_heading(doc, text, level=1, font_size = 14, italic=False, bold=False, indent=40):


    # 添加标题
    heading = doc.add_heading(text, level=level)

    heading.paragraph_format.left_indent = Pt(indent)  # 设置缩进 


    # 获取标题中的运行
    run = heading.runs[0]
    run.font.size = Pt(font_size)  # 设置字体大小
    run.bold = bold  # 设置加粗
    run.italic = italic  # 设置斜体

    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()  # 获取或添加 rPr 元素
    rFonts = rPr.get_or_add_rFonts()  # 获取或添加 rFonts 元素
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')  # 设置东亚字体为微軟正黑體
    rFonts.set(qn('w:ascii'), '微軟正黑體')  # 设置拉丁字母字体
    rFonts.set(qn('w:hAnsi'), '微軟正黑體')  # 设置 ANSI 字体
    rFonts.set(qn('w:cs'), '微軟正黑體')  # 设置其他语言字体


def add_custom_paragraph(doc, text, font_size=12, east_asia_font='微軟正黑體', indent=70):
    # 添加段落
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(indent)  # 设置左缩进


    # 添加运行并设置字体
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)  # 设置字体大小

    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()  # 获取或添加 rPr 元素
    rFonts = rPr.get_or_add_rFonts()  # 获取或添加 rFonts 元素
    rFonts.set(qn('w:eastAsia'), east_asia_font)  # 设置东亚字体    
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')  # 设置东亚字体为微軟正黑體
    rFonts.set(qn('w:ascii'), '微軟正黑體')  # 设置拉丁字母字体
    rFonts.set(qn('w:hAnsi'), '微軟正黑體')  # 设置 ANSI 字体
    rFonts.set(qn('w:cs'), '微軟正黑體')  # 设置其他语言字体



def risk_str_translate(risk_level):
    switcher = {
        "Critical": "重大風險",
        "High": "高風險",
        "Medium": "中風險",
        "Low": "低風險"
    }
    return switcher.get(risk_level, "無效的選擇")  # 默認值




# 创建新的 Word 文档
doc = Document()





# 1. 從檔案中讀取 JSON 數據
with open('output_2.json', 'r', encoding='utf-8') as file:
    data = json.load(file)

# 2. 在字典中遍歷所有鍵，並新增鍵值
for risk_level, items in data.items():


    risk_level_title = risk_str_translate(risk_level) 

    add_custom_heading(doc, risk_level_title, level=2, font_size = 16, indent=0)


    for index, item in enumerate(items):
        # 使用函数添加大标题和小标题
        add_custom_heading(doc, item["name"], level=3, indent=0)

        add_custom_heading(doc, '主機IP及服務相對應之埠號', level=4)
        for ip in item["address"]:
            add_custom_paragraph(doc, ip)

        add_custom_heading(doc, 'CVSS分數', level=4)
        add_custom_paragraph(doc, item["cvss"])

        add_custom_heading(doc, '風險敘述', level=4)
        add_custom_paragraph(doc, item["description"])

        add_custom_heading(doc, '修補建議', level=4)
        add_custom_paragraph(doc, item["solution"])


        # 判斷最後一圈且風險等級為low 不加上分頁
        if index == len(items) - 1 and risk_level == "Low":
            break  # 直接跳出迴圈

        # 在這裡添加分頁符
        doc.add_page_break()    




# 保存文档
doc.save('輸出近乎完美word_3.docx')
